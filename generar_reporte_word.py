import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def format_currency(val):
    if pd.isna(val):
        return "0"
    if val >= 1000:
        return f"{val/1000:.0f}K"
    return f"{val:.0f}"

def format_pct(val):
    if pd.isna(val):
        return "0%"
    return f"{val*100:.0f}%"

def set_cell_bg(cell, color_hex):
    # Apply background color to a cell
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
    cell._tc.get_or_add_tcPr().append(shading_elm)

def main():
    print("Cargando datos...")
    
    # 1. Ranking
    df_rank = pd.read_csv('webapp/Reportes_Individuales_CSV/Ranking.csv')
    df_rank['Is_Top300'] = df_rank['Ranking'] <= 300
    
    # 2. TRX Semanal
    df_trx = pd.read_csv('webapp/Reportes_Individuales_CSV/TRX_SEM_clean.csv')
    week_cols = [c for c in df_trx.columns if c.startswith('Semana ')]
    if len(week_cols) < 4:
        raise ValueError("No hay suficientes semanas de datos")
        
    last_week = week_cols[-1]
    prev_week = week_cols[-2]
    last_4_weeks = week_cols[-4:]
    
    # Convertir a numerico
    for col in last_4_weeks:
        df_trx[col] = pd.to_numeric(df_trx[col], errors='coerce').fillna(0)
        
    df_trx['Abonos'] = df_trx[last_week]
    df_trx['Prev'] = df_trx[prev_week]
    df_trx['Prom'] = df_trx[last_4_weeks].mean(axis=1)
    
    df_trx['Var_Pct'] = np.where(
        df_trx['Prev'] > 0,
        (df_trx['Abonos'] - df_trx['Prev']) / df_trx['Prev'],
        np.where(df_trx['Abonos'] < df_trx['Prev'], -1.0, 0.0) # Si prev=0 y abonos<0 (raro), -100%
    )
    
    # 3. Semaforo
    df_sem = pd.read_csv('webapp/Reportes_Individuales_CSV/Semaforo_Salud.csv')
    
    # 4. TPV vs SPEI
    df_wuzi = pd.read_csv('webapp/Reportes_Individuales_CSV/Wuzi_Semanal.csv')
    df_bp = pd.read_csv('webapp/Reportes_Individuales_CSV/BP_Semanal.csv')
    df_spei = pd.read_csv('webapp/Reportes_Individuales_CSV/SPEI_Semanal.csv')
    
    # Sumar totales
    wuzi_cols = [c for c in df_wuzi.columns if c != 'Holder']
    for c in wuzi_cols: df_wuzi[c] = pd.to_numeric(df_wuzi[c], errors='coerce').fillna(0)
    df_wuzi['Tot_Wuzi'] = df_wuzi[wuzi_cols].sum(axis=1)
    
    bp_cols = [c for c in df_bp.columns if c != 'Holder']
    for c in bp_cols: df_bp[c] = pd.to_numeric(df_bp[c], errors='coerce').fillna(0)
    df_bp['Tot_BP'] = df_bp[bp_cols].sum(axis=1)
    
    spei_cols = [c for c in df_spei.columns if c != 'Holder']
    for c in spei_cols: df_spei[c] = pd.to_numeric(df_spei[c], errors='coerce').fillna(0)
    df_spei['Tot_SPEI'] = df_spei[spei_cols].sum(axis=1)
    
    # Merge todo
    df = df_rank[['Holder', 'Vendedor', 'Is_Top300']].merge(df_trx[['Holder', 'Abonos', 'Prev', 'Prom', 'Var_Pct']], on='Holder', how='left')
    df = df.merge(df_sem[['Holder', 'Semáforo']], on='Holder', how='left')
    
    df = df.merge(df_wuzi[['Holder', 'Tot_Wuzi']], on='Holder', how='left')
    df = df.merge(df_bp[['Holder', 'Tot_BP']], on='Holder', how='left')
    df = df.merge(df_spei[['Holder', 'Tot_SPEI']], on='Holder', how='left')
    
    df['Tot_TPV'] = df['Tot_Wuzi'].fillna(0) + df['Tot_BP'].fillna(0)
    df['Tot_SPEI'] = df['Tot_SPEI'].fillna(0)
    df['Tot_General'] = df['Tot_TPV'] + df['Tot_SPEI']
    
    def get_tipo(row):
        tot = row['Tot_General']
        if tot == 0:
            return "Ninguno"
        if row['Tot_TPV'] / tot > 0.9:
            return "TPV"
        if row['Tot_SPEI'] / tot > 0.9:
            return "SPEI"
        return "Ambos"
        
    df['TPV/SPEI'] = df.apply(get_tipo, axis=1)
    
    # Ya no mapeamos el Semáforo de Salud, sino que la Salud dependerá de la Observación
    
    # Filtrar solo caidas
    # Una caida es Var_Pct < 0 y Prev > 0
    df_caidas = df[(df['Var_Pct'] < 0) & (df['Prev'] > 0)].copy()
    
    # 5. Filtrar por Vendedores excluidos
    excluded_vendedores = ['wuzi', 'paywise', 'nicole', 'sin vendedor']
    # Consideramos NaNs o vacíos como 'sin vendedor'
    vend_lower = df_caidas['Vendedor'].fillna('sin vendedor').str.strip().str.lower()
    df_caidas = df_caidas[~vend_lower.isin(excluded_vendedores)]
    
    # 6. Excluir los que ya tienen contexto
    try:
        df_context = pd.read_csv('webapp/Reportes_Individuales_CSV/Contexto_Holders.csv', encoding='utf-8-sig')
        # Merge de contexto
        if 'Holder' in df_context.columns and 'Contexto' in df_context.columns:
            df_caidas = df_caidas.merge(df_context[['Holder', 'Contexto']], on='Holder', how='left')
            
            # --- LÓGICA PARA EXCEL DE SEGUIMIENTO ---
            mask_con_contexto = (
                pd.notna(df_caidas['Contexto']) & 
                (df_caidas['Contexto'].str.strip() != '') & 
                (df_caidas['Contexto'].str.strip() != 'Sin contexto')
            )
            df_seguimiento = df_caidas[mask_con_contexto].copy()
            
            vendedores_seguimiento = [
                'Adam Dorenbaum', 'Alejandro Aguilera', 'Arturo Liogón', 'Iñaki Otegui',
                'Javier Lindner', 'Phil Caire', 'Santiago García', 'Alfredo Maccise',
                'Ana Maria Meza', 'Andre Joloy', 'Hector Hupy', 'Jorge Emilio',
                'Luis Frausto', 'Mónica Segovia', 'Ninel Chavarría', 'Rodrigo Siqueff'
            ]
            vendedores_lower = [v.lower() for v in vendedores_seguimiento]
            df_seguimiento = df_seguimiento[df_seguimiento['Vendedor'].fillna('').str.strip().str.lower().isin(vendedores_lower)]
            
            if not df_seguimiento.empty:
                # Format some columns if needed, or keep raw
                cols_excel = ['Vendedor', 'Holder', 'TPV/SPEI', 'Abonos', 'Prom', 'Var_Pct', 'Contexto']
                df_export = df_seguimiento[cols_excel].copy()
                df_export.to_excel('Seguimiento_Contextos.xlsx', index=False)
                print(f"Excel generado exitosamente con {len(df_export)} registros: Seguimiento_Contextos.xlsx")
            else:
                print("No se generó Excel de Seguimiento porque no hay registros que cumplan las condiciones.")
            # -----------------------------------------
            
            # Dejar solo los que no tienen contexto (NaN o vacío o 'Sin contexto') para el reporte de Word
            df_caidas = df_caidas[~mask_con_contexto]
    except Exception as e:
        print(f"No se pudo cargar o procesar Contexto_Holders.csv: {e}")
        
    # Observacion
    def get_obs(row):
        if row['Var_Pct'] <= -0.99:
            return "Riesgo abandono"
        elif row['Var_Pct'] <= -0.20:
            return "Caída relevante"
        else:
            return "Debajo de promedio"
            
    df_caidas['Observación'] = df_caidas.apply(get_obs, axis=1)
    
    # Mapear Salud a la Observacion
    def get_emoji(val):
        if val == "Riesgo abandono": return "🚨"
        elif val == "Caída relevante": return "⚠️"
        else: return "🟡"
        
    df_caidas['Salud'] = df_caidas['Observación'].apply(get_emoji)
    
    # Sort by vendor then drop severity
    df_caidas = df_caidas.sort_values(by=['Vendedor', 'Var_Pct'], ascending=[True, True])
    
    print(f"Encontradas {len(df_caidas)} caídas relevantes después de filtros.")
    
    # Crear Word
    doc = Document()
    doc.add_heading('📊 TRX Semana vs Semana – Caídas relevantes', 0)
    
    if len(df_caidas) == 0:
        doc.add_paragraph("No se registraron caídas relevantes con los filtros actuales.")
    else:
        # Group by Vendedor
        grouped = df_caidas.groupby('Vendedor')
        
        for vendedor, group in grouped:
            # Header for Vendedor
            vend_name = str(vendedor) if pd.notna(vendedor) else "Vendedor No Asignado"
            doc.add_heading(f'Vendedor: {vend_name}', level=2)
            
            # Helper to create tables
            def create_table_for_group(subgroup, title):
                doc.add_heading(title, level=3)
                if len(subgroup) == 0:
                    doc.add_paragraph("Sin caídas en esta categoría.")
                    return
                
                table = doc.add_table(rows=1, cols=7)
                table.style = 'Table Grid'
                
                headers = ['Cuenta', 'TPV/SPEI', 'Abonos', 'Prom', 'Var.', 'Salud', 'Observación']
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                    run = hdr_cells[i].paragraphs[0].runs[0]
                    run.font.bold = True
                    set_cell_bg(hdr_cells[i], "D9D9D9")
                    hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                for _, row in subgroup.iterrows():
                    row_cells = table.add_row().cells
                    
                    cuenta = str(row['Holder'])
                    row_cells[0].text = cuenta
                    
                    row_cells[1].text = row['TPV/SPEI']
                    row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    row_cells[2].text = format_currency(row['Abonos'])
                    row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    
                    row_cells[3].text = format_currency(row['Prom'])
                    row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    
                    row_cells[4].text = format_pct(row['Var_Pct'])
                    row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    
                    row_cells[5].text = row['Salud']
                    row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if len(row_cells[5].paragraphs[0].runs) > 0:
                        run = row_cells[5].paragraphs[0].runs[0]
                        run.font.name = 'Segoe UI Emoji'
                    
                    row_cells[6].text = row['Observación']
                
                doc.add_paragraph("")
            
            # Split Top 300 and No Top 300
            top300_group = group[group['Is_Top300'] == True].head(10)
            no_top300_group = group[group['Is_Top300'] == False].head(10)
            
            create_table_for_group(top300_group, "Top 300")
            create_table_for_group(no_top300_group, "No Top 300")

    output_path = 'Caidas_Relevantes_v6.docx'
    doc.save(output_path)
    print(f"Reporte generado exitosamente en: {output_path}")

if __name__ == '__main__':
    main()
