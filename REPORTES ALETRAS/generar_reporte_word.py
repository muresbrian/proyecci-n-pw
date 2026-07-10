import pandas as pd
import os
import glob
from docx import Document
from datetime import datetime, timedelta
import locale

def find_latest_report(directory):
    # Buscar primero los reportes que terminen en _Analizado.xlsx
    files = glob.glob(os.path.join(directory, 'Reporte_Alertas_*_Analizado.xlsx'))
    if not files:
        files = glob.glob(os.path.join(directory, 'Reporte_Alertas_*.xlsx'))
    if not files:
        return None
    latest_file = max(files, key=os.path.getmtime)
    return latest_file

def build_report():
    try:
        locale.setlocale(locale.LC_TIME, 'es_ES.UTF-8')
    except:
        pass

    # Obtener el directorio base de forma relativa (un nivel arriba de este script)
    report_dir = os.path.dirname(os.path.abspath(__file__))
    base_dir = os.path.dirname(report_dir)
    
    latest_report = find_latest_report(report_dir)
    if not latest_report:
        print("No se encontraron reportes descargados en la carpeta REPORTES ALETRAS.")
        return
    
    print(f"Utilizando el reporte: {os.path.basename(latest_report)}")
    
    try:
        df_alertas = pd.read_excel(latest_report, sheet_name='Alertas Completas')
    except ValueError:
        df_alertas = pd.read_excel(latest_report, sheet_name=0)
    
    context_path = os.path.join(base_dir, "webapp", "Reportes_Individuales_CSV", "Contexto_Holders.csv")
    if os.path.exists(context_path):
        df_context = pd.read_csv(context_path, encoding='utf-8-sig')
    else:
        df_context = pd.DataFrame(columns=['Holder', 'Contexto'])
        
    tipo_path = os.path.join(report_dir, "TIPO_TRX.xlsx")
    if os.path.exists(tipo_path):
        df_tipo = pd.read_excel(tipo_path)
        if 'Tipo' in df_tipo.columns:
            df_tipo = df_tipo.rename(columns={'Tipo': 'TPV / SPEI'})
    else:
        df_tipo = pd.DataFrame(columns=['Holder', 'TPV / SPEI'])
        
    df_merged = df_alertas.copy()
    
    if 'Holder' in df_tipo.columns and 'TPV / SPEI' in df_tipo.columns:
        df_merged['_merge_key_tipo'] = df_merged['Holder'].astype(str).str.strip().str.upper()
        df_tipo['_merge_key_tipo'] = df_tipo['Holder'].astype(str).str.strip().str.upper()
        df_tipo_dedup = df_tipo[['_merge_key_tipo', 'TPV / SPEI']].drop_duplicates(subset=['_merge_key_tipo'])
        df_merged = df_merged.merge(df_tipo_dedup, on='_merge_key_tipo', how='left')
        df_merged = df_merged.drop(columns=['_merge_key_tipo'])
    else:
        df_merged['TPV / SPEI'] = 'Desconocido'
        
    df_merged['TPV / SPEI'] = df_merged['TPV / SPEI'].fillna('Desconocido')
    
    if 'Holder' in df_context.columns and 'Contexto' in df_context.columns:
        df_merged['_merge_key_ctx'] = df_merged['Holder'].astype(str).str.strip().str.upper()
        df_context['_merge_key_ctx'] = df_context['Holder'].astype(str).str.strip().str.upper()
        df_context_dedup = df_context[['_merge_key_ctx', 'Contexto']].drop_duplicates(subset=['_merge_key_ctx'])
        df_merged = df_merged.merge(df_context_dedup, on='_merge_key_ctx', how='left')
        df_merged = df_merged.drop(columns=['_merge_key_ctx'])
    else:
        df_merged['Contexto'] = 'Sin contexto'
        
    df_merged['Contexto'] = df_merged['Contexto'].fillna('Sin contexto')
    
    if 'Vendedor' not in df_merged.columns:
        df_merged['Vendedor'] = 'Desconocido'
        
    # Excluir a los vendedores Paywise, Wuzi, Nicole Bauzá y cuentas Sin Vendedor (sin importar mayúsculas o minúsculas)
    vendedores_excluidos = ['paywise', 'wuzi', 'pay wise', 'nicole bauza', 'nicole bauzá', 'sin vendedor']
    df_merged = df_merged[~df_merged['Vendedor'].str.lower().isin(vendedores_excluidos)]
    
    # Identificar columna Falta_Normal (Columna I en excel analizado)
    falta_normal_col = None
    if 'Falta_Normal' in df_merged.columns:
        falta_normal_col = 'Falta_Normal'
    elif len(df_merged.columns) > 8 and df_merged.columns[8] not in ['Alerta', 'Tipo_alerta']:
        falta_normal_col = df_merged.columns[8]
        
    if falta_normal_col:
        df_merged['Falta Normal'] = df_merged[falta_normal_col].fillna('No aplica')
    else:
        df_merged['Falta Normal'] = 'Normal'
        
    # Simplificar la columna Falta Normal para cuentas con más de 7 días (para evitar romper la tabla)
    if 'Dias_sin_transaccionar' in df_merged.columns:
        def clean_val(val):
            val_str = str(val).lower()
            if 'inusual' in val_str:
                return 'No, inusual'
            if 'normal' in val_str:
                return 'Normal'
            return str(val)
            
        df_merged.loc[df_merged['Dias_sin_transaccionar'] > 7, 'Falta Normal'] = df_merged.loc[
            df_merged['Dias_sin_transaccionar'] > 7, 'Falta Normal'
        ].apply(clean_val)
    
    # Extraer la fecha del archivo de reporte para usarla como ancla de probabilidad
    nombre_base = os.path.basename(latest_report)
    fecha_str = nombre_base.replace('Reporte_Alertas_', '').replace('_Analizado', '').replace('.xlsx', '').split(' ')[0]
    try:
        fecha_reporte = datetime.strptime(fecha_str, '%Y%m%d')
    except:
        fecha_reporte = datetime.now()

    # Calcular Probabilidad de No Transaccionar
    nombres_dias_col = {
        0: '%_Lunes', 1: '%_Martes', 2: '%_Miercoles', 3: '%_Jueves', 4: '%_Viernes', 5: '%_Sabado', 6: '%_Domingo'
    }
    nombres_dias_corto = {
        0: 'Lun', 1: 'Mar', 2: 'Mie', 3: 'Jue', 4: 'Vie', 5: 'Sab', 6: 'Dom'
    }
    
    def calculate_prob_no_trx(row):
        dias_sin_trx = row.get('Dias_sin_transaccionar', 0)
        if pd.isna(dias_sin_trx) or dias_sin_trx <= 0:
            return 'Normal'
            
        dias_sin_trx = int(dias_sin_trx)
        if dias_sin_trx > 7:
            return 'Alta'
            
        probs = []
        for i in range(1, dias_sin_trx + 1):
            dia_evaluado = fecha_reporte - timedelta(days=i-1)
            dia_semana = dia_evaluado.weekday()
            col_name = nombres_dias_col[dia_semana]
            
            if col_name in row and not pd.isna(row[col_name]):
                frecuencia_trx = row[col_name]
                prob_no_trx = max(0, min(100, 100 - frecuencia_trx))
                probs.append(f"{nombres_dias_corto[dia_semana]}: {prob_no_trx:.0f}%")
            else:
                probs.append(f"{nombres_dias_corto[dia_semana]}: 100%")
        return ", ".join(probs)
        
    df_merged['Probabilidad No TRX'] = df_merged.apply(calculate_prob_no_trx, axis=1)

    df_merged = df_merged.rename(columns={'Holder': 'Cuenta'})
    
    if 'Alerta' not in df_merged.columns and 'Tipo_alerta' in df_merged.columns:
        df_merged = df_merged.rename(columns={'Tipo_alerta': 'Alerta'})
        
    critico_mask = df_merged['Alerta'] == 'Más de 7 días sin TRX'
    riesgo_mask = df_merged['Alerta'] == '4 a 7 días sin Trx'
    alerta_3d_mask = df_merged['Alerta'] == '3 días seguidos sin Trx'
    alerta_2d_mask = df_merged['Alerta'] == '2 días seguidos sin Trx'
    alerta_1d_mask = df_merged['Alerta'] == '1 día sin Trx'
    
    groups = [
        ("🔴 +7 días sin TRX (CRÍTICO)", df_merged[critico_mask]),
        ("🔶 4–7 días sin TRX (RIESGO)", df_merged[riesgo_mask]),
        ("🟡 3 días sin TRX (ALERTA)", df_merged[alerta_3d_mask]),
        ("🟡 2 días sin TRX (ALERTA)", df_merged[alerta_2d_mask]),
        ("🟡 1 día sin TRX (ALERTA)", df_merged[alerta_1d_mask])
    ]
    
    months = ["Ene", "Feb", "Mar", "Abr", "May", "Jun", "Jul", "Ago", "Sep", "Oct", "Nov", "Dic"]
    now = datetime.now()
    date_str = f"{fecha_reporte.day} {months[fecha_reporte.month-1]} {fecha_reporte.year}"
    
    for show_prob in [True, False]:
        doc = Document()
        doc.add_heading(f"📊 TRX Diario – {date_str}", 0)
        
        for title, df_group in groups:
            total = len(df_group)
            p = doc.add_paragraph()
            runner = p.add_run(f"{title} | Total: {total}")
            runner.bold = True
            
            if total == 0:
                doc.add_paragraph("No hay cuentas en esta categoría.")
                doc.add_paragraph("")
                continue
                
            df_group = df_group.sort_values(['Vendedor', 'Cuenta'])
            
            for vendedor, df_vend in df_group.groupby('Vendedor', sort=False):
                # Agregar el nombre del vendedor antes de la tabla
                p_vend = doc.add_paragraph()
                runner_vend = p_vend.add_run(f"Vendedor: {vendedor}")
                runner_vend.bold = True
                
                cols_count = 5 if show_prob else 4
                table = doc.add_table(rows=1, cols=cols_count)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Cuenta'
                hdr_cells[1].text = 'TPV / SPEI'
                hdr_cells[2].text = 'Comportamiento'
                if show_prob:
                    hdr_cells[3].text = 'Prob. No TRX'
                    hdr_cells[4].text = 'Contexto'
                else:
                    hdr_cells[3].text = 'Contexto'
                
                for cell in hdr_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                for _, row in df_vend.iterrows():
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(row['Cuenta'])
                    row_cells[1].text = str(row['TPV / SPEI'])
                    row_cells[2].text = str(row['Falta Normal'])
                    if show_prob:
                        row_cells[3].text = str(row['Probabilidad No TRX'])
                        row_cells[4].text = str(row['Contexto'])
                    else:
                        row_cells[3].text = str(row['Contexto'])
                    
                doc.add_paragraph("")
                
        # Guardar con nombres distintos según la versión
        suffix = "" if show_prob else "_Sin_Frecuencia"
        output_filename = os.path.join(report_dir, f"Reporte_Ejecutivo{suffix}_{fecha_reporte.strftime('%Y%m%d')}_{now.strftime('%H%M%S')}.docx")
        doc.save(output_filename)
        print(f"¡Éxito! Reporte Word generado: {output_filename}")

if __name__ == "__main__":
    build_report()
